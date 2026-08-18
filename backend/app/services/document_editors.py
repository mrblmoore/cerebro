"""
Editing documents: Word and Excel.

Edits are described as a list of operations rather than a rewritten file, which
keeps three properties that matter when software changes someone's work:

* **Reviewable** — every operation says exactly what it will do, and ``dry_run``
  reports that without touching the file.
* **Reversible** — the original is copied to a timestamped backup first.
* **Safe to refuse** — if Word or Excel has the file open, the edit is rejected
  rather than racing the application and losing one side of the change.
"""

import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from app.core import logger
from app.core.config import settings
from app.services.document_readers import DocumentError, classify, is_open_in_office

#: "007", "00123" — an identifier that happens to be digits, not a number.
_LEADING_ZERO = re.compile(r"^-?0\d")

#: Operations, by document kind. Each maps to a handler below.
OPERATIONS = {
    "docx": {
        "replace_text": "Replace every occurrence of some text",
        "append_paragraph": "Add a paragraph at the end",
        "insert_paragraph": "Insert a paragraph at a position",
        "set_paragraph": "Rewrite one paragraph",
        "delete_paragraph": "Remove one paragraph",
        "append_heading": "Add a heading at the end",
    },
    "xlsx": {
        "set_cell": "Set one cell's value or formula",
        "set_range": "Set a rectangular block of values",
        "append_row": "Add a row to the end of a sheet",
        "clear_cell": "Empty a cell",
        "add_sheet": "Add a new sheet",
        "rename_sheet": "Rename a sheet",
    },
}


def describe_operations() -> Dict[str, Any]:
    return {kind: [{"op": name, "description": text} for name, text in operations.items()]
            for kind, operations in OPERATIONS.items()}


def backup(path: Path) -> Optional[Path]:
    """Copy the file next to itself with a timestamp, before anything changes."""
    if not settings.DOCUMENT_BACKUP_ON_EDIT:
        return None
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    target = path.with_name(f"{path.stem}.cerebro-backup-{stamp}{path.suffix}")
    shutil.copy2(path, target)
    return target


def _discard_backup(path: Optional[Path]) -> None:
    """Remove a backup taken for an edit that never happened."""
    if path is None:
        return
    try:
        path.unlink(missing_ok=True)
    except OSError:
        pass


def _require_writable(path: Path) -> None:
    if not path.exists():
        raise DocumentError(f"File not found: {path}")
    if is_open_in_office(path):
        raise DocumentError(
            f"{path.name} is open in Office right now. Close it and try again — "
            "editing underneath a running Word or Excel loses one set of changes."
        )


# -------------------------------------------------------------------- Word
def _edit_docx(path: Path, operations: List[Dict[str, Any]],
               dry_run: bool) -> List[Dict[str, Any]]:
    try:
        import docx
    except ImportError as exc:
        raise DocumentError(
            "Editing Word files needs python-docx. "
            "pip install -r backend/requirements-documents.txt") from exc

    document = docx.Document(str(path))
    applied: List[Dict[str, Any]] = []

    for operation in operations:
        name = operation.get("op")

        if name == "replace_text":
            find = operation.get("find")
            replace = operation.get("replace", "")
            if not find:
                raise DocumentError("replace_text needs a 'find' value.")

            count = 0
            for paragraph in document.paragraphs:
                if find not in paragraph.text:
                    continue
                # Word splits text across runs arbitrarily; rewriting the first
                # run and clearing the rest preserves the paragraph's formatting
                # while guaranteeing the replacement actually happens.
                replaced = paragraph.text.replace(find, replace)
                count += paragraph.text.count(find)
                if not dry_run:
                    if paragraph.runs:
                        paragraph.runs[0].text = replaced
                        for run in paragraph.runs[1:]:
                            run.text = ""
                    else:
                        paragraph.text = replaced
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if find in cell.text:
                            count += cell.text.count(find)
                            if not dry_run:
                                cell.text = cell.text.replace(find, replace)
            applied.append({"op": name, "occurrences": count,
                            "detail": f"{count} occurrence(s) of {find!r} → {replace!r}"})

        elif name == "append_paragraph":
            text = operation.get("text", "")
            if not dry_run:
                document.add_paragraph(text, style=operation.get("style") or None)
            applied.append({"op": name, "detail": f"appended {text[:60]!r}"})

        elif name == "append_heading":
            text = operation.get("text", "")
            level = int(operation.get("level", 1))
            if not dry_run:
                document.add_heading(text, level=max(0, min(level, 9)))
            applied.append({"op": name, "detail": f"heading {level}: {text[:60]!r}"})

        elif name in ("insert_paragraph", "set_paragraph", "delete_paragraph"):
            index = operation.get("index")
            if index is None:
                raise DocumentError(f"{name} needs an 'index'.")
            index = int(index)
            if not 0 <= index < len(document.paragraphs):
                raise DocumentError(
                    f"Paragraph {index} does not exist — the document has "
                    f"{len(document.paragraphs)}.")

            paragraph = document.paragraphs[index]
            if name == "set_paragraph":
                text = operation.get("text", "")
                if not dry_run:
                    if paragraph.runs:
                        paragraph.runs[0].text = text
                        for run in paragraph.runs[1:]:
                            run.text = ""
                    else:
                        paragraph.text = text
                applied.append({"op": name, "detail": f"paragraph {index} → {text[:60]!r}"})
            elif name == "delete_paragraph":
                if not dry_run:
                    element = paragraph._element
                    element.getparent().remove(element)
                applied.append({"op": name, "detail": f"removed paragraph {index}"})
            else:
                text = operation.get("text", "")
                if not dry_run:
                    new = paragraph.insert_paragraph_before(text)
                    if operation.get("style"):
                        new.style = operation["style"]
                applied.append({"op": name, "detail": f"inserted before {index}"})

        else:
            raise DocumentError(
                f"Unknown Word operation {name!r}. Supported: "
                f"{', '.join(OPERATIONS['docx'])}")

    if not dry_run:
        document.save(str(path))
    return applied


# ------------------------------------------------------------------- Excel
def _sheet(workbook, name: Optional[str]):
    if not name:
        return workbook.active
    if name not in workbook.sheetnames:
        raise DocumentError(
            f"No sheet named {name!r}. This workbook has: {', '.join(workbook.sheetnames)}")
    return workbook[name]


def _coerce(value: Any) -> Any:
    """
    Let plain numbers land as numbers so Excel can compute with them.

    Deliberately conservative about what counts as a number: leading zeros and a
    leading ``+`` mean the string is an identifier, not a quantity. Coercing
    ``"00123"`` to ``123`` or ``"+441234567"`` to an integer silently corrupts
    case numbers, part codes and phone numbers written into a spreadsheet.
    """
    if not isinstance(value, str):
        return value

    text = value.strip()
    if not text or text.startswith("="):
        return value
    if text.startswith("+"):
        return value
    if _LEADING_ZERO.match(text):
        return value

    try:
        return int(text)
    except ValueError:
        pass
    try:
        return float(text)
    except ValueError:
        return value


def _edit_xlsx(path: Path, operations: List[Dict[str, Any]],
               dry_run: bool) -> List[Dict[str, Any]]:
    try:
        import openpyxl
    except ImportError as exc:
        raise DocumentError(
            "Editing Excel files needs openpyxl. "
            "pip install -r backend/requirements-documents.txt") from exc

    workbook = openpyxl.load_workbook(str(path))
    applied: List[Dict[str, Any]] = []

    try:
        for operation in operations:
            name = operation.get("op")

            if name in ("set_cell", "clear_cell"):
                cell_ref = operation.get("cell")
                if not cell_ref:
                    raise DocumentError(f"{name} needs a 'cell' like 'B7'.")
                worksheet = _sheet(workbook, operation.get("sheet"))
                value = None if name == "clear_cell" else _coerce(operation.get("value"))
                previous = worksheet[cell_ref].value
                if not dry_run:
                    worksheet[cell_ref] = value
                applied.append({
                    "op": name, "sheet": worksheet.title, "cell": cell_ref,
                    "detail": f"{worksheet.title}!{cell_ref}: {previous!r} → {value!r}",
                })

            elif name == "set_range":
                start = operation.get("start")
                values = operation.get("values")
                if not start or not isinstance(values, list):
                    raise DocumentError("set_range needs 'start' and a list of row lists.")
                worksheet = _sheet(workbook, operation.get("sheet"))
                origin = worksheet[start]
                changed = 0
                for row_offset, row in enumerate(values):
                    row = row if isinstance(row, list) else [row]
                    for column_offset, value in enumerate(row):
                        if not dry_run:
                            worksheet.cell(row=origin.row + row_offset,
                                           column=origin.column + column_offset,
                                           value=_coerce(value))
                        changed += 1
                applied.append({"op": name, "sheet": worksheet.title,
                                "detail": f"{changed} cell(s) from {start}"})

            elif name == "append_row":
                values = operation.get("values")
                if not isinstance(values, list):
                    raise DocumentError("append_row needs a 'values' list.")
                worksheet = _sheet(workbook, operation.get("sheet"))
                if not dry_run:
                    worksheet.append([_coerce(v) for v in values])
                applied.append({"op": name, "sheet": worksheet.title,
                                "detail": f"row {worksheet.max_row + (0 if dry_run else 0)} "
                                          f"← {values[:6]}"})

            elif name == "add_sheet":
                title = operation.get("title") or operation.get("name")
                if not title:
                    raise DocumentError("add_sheet needs a 'title'.")
                if title in workbook.sheetnames:
                    raise DocumentError(f"A sheet named {title!r} already exists.")
                if not dry_run:
                    workbook.create_sheet(title=title)
                applied.append({"op": name, "detail": f"created sheet {title!r}"})

            elif name == "rename_sheet":
                worksheet = _sheet(workbook, operation.get("sheet"))
                title = operation.get("title")
                if not title:
                    raise DocumentError("rename_sheet needs a 'title'.")
                previous = worksheet.title
                if not dry_run:
                    worksheet.title = title
                applied.append({"op": name, "detail": f"{previous!r} → {title!r}"})

            else:
                raise DocumentError(
                    f"Unknown Excel operation {name!r}. Supported: "
                    f"{', '.join(OPERATIONS['xlsx'])}")

        if not dry_run:
            workbook.save(str(path))
    finally:
        workbook.close()

    return applied


EDITORS = {"docx": _edit_docx, "xlsx": _edit_xlsx}


def apply(path: Path, operations: List[Dict[str, Any]],
          dry_run: bool = False) -> Dict[str, Any]:
    """
    Apply operations to a document.

    With ``dry_run`` the file is opened and every operation validated against its
    real content — an out-of-range paragraph or a misspelled sheet name fails
    here — but nothing is written. That makes "show me what you would change"
    genuinely trustworthy rather than a guess.
    """
    path = Path(path)
    _require_writable(path)

    kind = classify(path)
    if kind not in EDITORS:
        raise DocumentError(
            f"Cerebro can read {kind or 'this file'} but not edit it. "
            f"Editable formats: Word (.docx) and Excel (.xlsx)."
        )
    if not operations:
        raise DocumentError("No operations given — nothing to do.")

    saved_backup = None
    if not dry_run:
        saved_backup = backup(path)

    try:
        applied = EDITORS[kind](path, operations, dry_run)
    except DocumentError:
        # Validation failed before anything was written, so the backup is a copy
        # of an unchanged file. Leaving it behind litters the user's folder —
        # and OneDrive would dutifully sync the litter.
        _discard_backup(saved_backup)
        raise
    except Exception as exc:
        if saved_backup and saved_backup.exists():
            shutil.copy2(saved_backup, path)
            raise DocumentError(
                f"Edit failed and the file was restored from backup: {exc}") from exc
        raise DocumentError(f"Edit failed: {exc}") from exc

    if not dry_run:
        logger.info("documents", "Edited document", {
            "path": str(path), "operations": len(applied),
            "backup": str(saved_backup) if saved_backup else None,
        })

    return {
        "ok": True,
        "path": str(path),
        "kind": kind,
        "dry_run": dry_run,
        "operations": applied,
        "backup": str(saved_backup) if saved_backup else None,
    }
