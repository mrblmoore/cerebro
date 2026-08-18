#!/usr/bin/env python3
"""
Cerebro test suite.

Runs against a throwaway SQLite database and needs no external services, no API
key and no running server:

    python test_mvp.py
"""

import os
import shutil
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT / "backend"))

# Point every component at a temporary database before the app imports settings.
_TEMP_DIR = tempfile.mkdtemp(prefix="cerebro-test-")
os.environ["DATABASE_URL"] = f"sqlite:///{Path(_TEMP_DIR).as_posix()}/test.db"
os.environ["LOG_TO_STDOUT"] = "false"
os.environ["CEREBRO_LOG_PATH"] = str(Path(_TEMP_DIR) / "test.log")
os.environ["LLM_PROVIDER"] = "none"
os.environ["VECTOR_BACKEND"] = "local"
os.environ["ENTERPRISE_INBOX_DIR"] = str(Path(_TEMP_DIR) / "inbox")
os.environ["ENTERPRISE_OUTBOX_DIR"] = str(Path(_TEMP_DIR) / "outbox")
os.environ["ENTERPRISE_ENABLED"] = "true"
os.environ["SHAREPOINT_SYNC_ROOTS"] = str(Path(_TEMP_DIR) / "sync")
os.environ["MEMORY_ENABLED"] = "true"
os.environ["STYLE_LEARNING_ENABLED"] = "true"
os.environ["TASKS_ENABLED"] = "true"
os.environ["NUDGES_ENABLED"] = "true"
os.environ["COPILOT_BRIDGE_ENABLED"] = "true"
os.environ["COPILOT_BRIDGE_DIR"] = str(Path(_TEMP_DIR) / "copilot")

from app.core.database import SessionLocal, init_db  # noqa: E402
from app.schemas.event import EventCreate  # noqa: E402
from app.services import embeddings  # noqa: E402
from app.services.context_engine import ContextEngine  # noqa: E402
from app.services.event_detector import EventDetector  # noqa: E402
from app.services.llm_service import LLMService  # noqa: E402
from app.services.rag_service import RAGService  # noqa: E402

PASSED = []
FAILED = []


def check(label, condition):
    (PASSED if condition else FAILED).append(label)
    print(f"  {'✓' if condition else '✗'} {label}")


def session():
    return SessionLocal()


# --------------------------------------------------------------- detectors
def test_event_detector():
    print("\nEvent detection")

    detected = EventDetector.detect_crm_event(
        "https://company.lightning.force.com/case/500abc123",
        "Case #12345 - Outlook Issue | John Doe",
    )
    check("Salesforce case detected", detected and detected[0] == "CRM_CASE_OPENED")
    check("Salesforce case id extracted", detected and detected[1]["case_id"] == "500abc123")

    detected = EventDetector.detect_remote_session_event(
        "bomgar-rep.exe", "BeyondTrust Remote Support - Connected to SERVER-01")
    check("Bomgar session detected", detected and detected[0] == "REMOTE_SESSION_CONNECTED")
    check("Remote host extracted", detected and detected[1]["host"] == "SERVER-01")

    detected = EventDetector.detect_remote_session_event(
        "bomgar-rep.exe", "BeyondTrust Remote Support - Session ended")
    check("Disconnection detected",
          detected and detected[0] == "REMOTE_SESSION_DISCONNECTED")

    detected = EventDetector.detect_crm_event(
        "https://acme.service-now.com/nav_to.do?uri=incident.do%3Fsysparm_query=number%3DINC0012345",
        "Northwind | ServiceNow")
    check("ServiceNow incident detected", detected and detected[1]["case_id"] == "INC0012345")

    detected = EventDetector.detect_call_event("Teams.exe", "Meeting with Contoso | Microsoft Teams")
    check("Call start detected", detected and detected[0] == "CALL_STARTED")

    detected = EventDetector.detect_call_event("Teams.exe", "Call ended | Microsoft Teams")
    check("Call end detected", detected and detected[0] == "CALL_ENDED")

    check("Non-CRM page ignored",
          EventDetector.detect_crm_event("https://news.example.com/", "Example") is None)
    check("Non-conferencing app ignored",
          EventDetector.detect_call_event("notepad.exe", "Untitled - Notepad") is None)


# ----------------------------------------------------------- context engine
def test_context_engine():
    print("\nContext engine")
    db = session()
    engine = ContextEngine(db)

    check("Context initialises", engine.init_context() is not None)

    result = engine.process_event(EventCreate(
        event_type="CRM_CASE_OPENED", source="test",
        data={"system": "Salesforce", "case_id": "12345", "customer": "Acme Corp"}))
    check("Case opened sets case", result["context"]["crm_case"] == "12345")
    check("Case opened sets customer", result["context"]["customer"] == "Acme Corp")
    check("Case opened suggests documentation",
          any(r["type"] == "retrieve_docs" for r in result["recommendations"]))

    result = engine.process_event(EventCreate(
        event_type="CALL_STARTED", source="test", data={"application": "Teams"}))
    check("Call started sets call_active", result["context"]["call_active"] is True)

    result = engine.process_event(EventCreate(
        event_type="REMOTE_SESSION_CONNECTED", source="test", data={"host": "SERVER-01"}))
    check("Remote session tracked", result["context"]["remote_session_active"] is True)
    check("Remote host recorded", result["context"]["remote_host"] == "SERVER-01")

    result = engine.process_event(EventCreate(
        event_type="UNKNOWN_EVENT_TYPE", source="test", data={}))
    check("Unknown event does not raise", result["event_id"] is not None)

    context = engine.reset_context()
    check("Reset clears the case", context.crm_case is None)
    check("Reset clears session state",
          context.call_active is False and context.remote_session_active is False)

    db.close()


def test_event_flow():
    print("\nFull event flow")
    db = session()
    engine = ContextEngine(db)
    engine.reset_context()

    for event_type, data in [
        ("CRM_CASE_OPENED", {"system": "Salesforce", "case_id": "500abc", "customer": "Contoso"}),
        ("CALL_STARTED", {"application": "Teams"}),
        ("REMOTE_SESSION_CONNECTED", {"host": "CONTOSO-PC"}),
        ("CALL_ENDED", {}),
    ]:
        engine.process_event(EventCreate(event_type=event_type, source="test", data=data))

    context = engine.get_current_context()
    check("Case survives the flow", context.crm_case == "500abc")
    check("Customer survives the flow", context.customer == "Contoso")
    check("Call ended", context.call_active is False)
    check("Remote session still open", context.remote_session_active is True)

    recommendations = engine.current_recommendations()
    check("Live recommendations produced", len(recommendations) > 0)
    check("Missing AI provider is surfaced",
          any(r["type"] == "configure_ai" for r in recommendations))

    db.close()


# -------------------------------------------------------------- embeddings
def test_embeddings():
    print("\nEmbeddings")
    related = embeddings.cosine(
        embeddings.local_embedding("Outlook cannot connect to the Exchange server"),
        embeddings.local_embedding("Exchange server unreachable from Outlook"))
    unrelated = embeddings.cosine(
        embeddings.local_embedding("Outlook cannot connect to the Exchange server"),
        embeddings.local_embedding("The printer queue is stuck and will not clear"))

    check(f"Related texts score higher ({related:.2f} > {unrelated:.2f})", related > unrelated)
    check("Identical text scores ~1.0", abs(embeddings.cosine(
        embeddings.local_embedding("same text"),
        embeddings.local_embedding("same text")) - 1.0) < 1e-6)
    check("Empty text is handled", embeddings.local_embedding("") is not None)


# --------------------------------------------------------- knowledge search
def test_knowledge_search():
    print("\nKnowledge search")
    db = session()
    rag = RAGService(db)
    check("Falls back to the built-in store", rag.backend == "local")

    rag.index_document({
        "title": "KB-1043 Outlook connectivity 0x80040115",
        "content": "Outlook cannot connect to Exchange. Error code 0x80040115 usually "
                   "means the Exchange server is unreachable. Check autodiscover and VPN.",
        "source": "RightAnswers",
    })
    rag.index_document({
        "title": "Printer spooler restart runbook",
        "content": "When a printer queue stalls, restart the print spooler service "
                   "and clear the spool folder.",
        "source": "runbook",
    })

    results = rag.search("outlook cannot reach exchange", limit=3)
    check("Search returns a result", len(results) > 0)
    check("Most relevant document ranks first",
          results and "Outlook" in results[0]["title"])

    results = rag.search("printer queue stuck", limit=3)
    check("Second query finds the other document",
          results and "Printer" in results[0]["title"])

    check("Empty query returns nothing", rag.search("", limit=3) == [])

    try:
        rag.index_document({"title": "Empty", "content": "   "})
        check("Empty content is rejected", False)
    except ValueError:
        check("Empty content is rejected", True)

    db.close()


# --------------------------------------------------------------------- llm
def test_llm_disabled():
    print("\nAI provider (disabled)")
    llm = LLMService()
    check("Reports itself disabled", llm.enabled is False)
    check("Status is not an error", llm.status()["ok"] is True)

    summary = llm.generate_case_summary({"customer": "Contoso", "title": "Outlook issue"})
    check("Generation returns guidance, not an exception", "Settings" in summary)


# ------------------------------------------------- enterprise bridge
def test_enterprise_normalisation():
    """Power Automate payloads vary; the normaliser must absorb the variation."""
    print("\nEnterprise bridge — normalising")
    from app.services.enterprise_service import assess_urgency, detect_case, normalise

    outlook = normalise({
        "source": "outlook", "type": "email",
        "timestamp": "2026-08-17T15:30:01Z",
        "sender": "person@company.com",
        "recipients": ["you@company.com"],
        "subject": "Need update on customer case",
        "body": "Can you send the latest status?",
        "thread_id": "abc123",
        "metadata": {"importance": "high"},
    })
    check("Outlook payload normalised", outlook["source"] == "outlook")
    check("Recipients flattened", outlook["recipients"] == "you@company.com")
    check("High importance raises urgency", outlook["urgency"] == "high")
    check("Timestamp parsed", outlook["timestamp"] is not None)

    teams = normalise({
        "source": "teams", "type": "message",
        "sender": "manager@company.com",
        "chat_or_channel": "Support Escalations",
        "body": "Can you check this? Case 500XY7 is escalating.",
        "thread_id": "teams-thread-456",
    })
    check("Teams channel kept", teams["chat_or_channel"] == "Support Escalations")
    check("Case found in the body", teams["case_id"] == "500XY7")
    check("Escalation raises urgency", teams["urgency"] == "high")

    # Graph-shaped fields, as a flow built from dynamic content produces them.
    graph = normalise({
        "source": "outlook",
        "from": {"emailAddress": {"address": "dana@contoso.com", "name": "Dana Reed"}},
        "toRecipients": [{"emailAddress": {"address": "you@company.com"}}],
        "subject": "RE: INC0012345",
        "bodyPreview": "<html><body><p>Any update?</p><b>Still down</b></body></html>",
        "receivedDateTime": "2026-08-17T16:02:00Z",
        "conversationId": "AAQkAD00",
    })
    check("Graph sender extracted", graph["sender"] == "dana@contoso.com")
    check("Graph display name kept", graph["sender_name"] == "Dana Reed")
    check("HTML stripped from the body", "<" not in (graph["body"] or ""))
    check("Body text survived stripping", "Any update?" in (graph["body"] or ""))
    check("Case found in the subject", graph["case_id"] == "INC0012345")

    check("Ordinary English is not a case", detect_case("In case of emergency") is None)
    check("Question raises urgency to medium",
          assess_urgency("Quick question", "Can you look at this?", "normal")[0] == "medium")
    check("Urgency always explains itself",
          bool(assess_urgency("URGENT", "outage", "normal")[1]))


def test_enterprise_ingest_and_reply():
    """Ingest is idempotent, and replies only leave on approval."""
    print("\nEnterprise bridge — ingest and reply")
    import json as _json

    from app.services import enterprise_service
    from app.services.enterprise_service import EnterpriseService

    inbox = Path(_TEMP_DIR) / "inbox"
    inbox.mkdir(parents=True, exist_ok=True)
    payload = {
        "source": "outlook", "type": "email", "external_id": "msg-1",
        "sender": "person@company.com", "subject": "Urgent: mail is down",
        "body": "Nothing is sending.", "timestamp": "2026-08-17T15:30:01Z",
    }
    (inbox / "outlook-1.json").write_text(_json.dumps(payload), encoding="utf-8")
    os.utime(inbox / "outlook-1.json", (0, 0))   # old enough to be swept

    db = session()
    result = enterprise_service.drain_inbox(db)
    check("File ingested", result["ingested"] == 1)
    check("File archived out of the inbox",
          not (inbox / "outlook-1.json").exists())

    # Replaying the same message must not create a second row.
    service = EnterpriseService(db)
    again = service.ingest_payload(payload, source_file="replay.json")
    check("Replay detected as duplicate", again["status"] == "duplicate")

    messages = service.list_messages()
    check("Exactly one message stored", len(messages) == 1)
    check("Urgency assessed on ingest", messages[0].urgency == "high")

    briefing = service.briefing(hours=24)
    check("Briefing counts the message", briefing["total"] == 1)
    check("Briefing lists it as urgent", len(briefing["urgent"]) == 1)

    outbox = Path(_TEMP_DIR) / "outbox"
    draft = service.create_action("reply_email", "On it — fix deploying now.",
                                  in_reply_to=messages[0].id, send=False)
    check("Reply starts as a draft", draft.status == "draft")
    check("Reply addressed to the sender", draft.to == "person@company.com")
    check("Subject prefixed with Re:", (draft.subject or "").startswith("Re:"))
    check("Nothing written to the outbox before approval",
          not outbox.exists() or not list(outbox.glob("*.json")))

    sent = service.dispatch_action(draft)
    check("Approval queues the reply", sent.status == "queued")

    files = list(outbox.glob("*.json"))
    check("Exactly one file written for Power Automate", len(files) == 1)
    written = _json.loads(files[0].read_text())
    check("Outbound payload carries the action", written["action"] == "reply_email")
    check("Outbound payload carries the recipient",
          written["to"] == ["person@company.com"])
    check("No partial .part file left behind", not list(outbox.glob("*.part")))
    db.close()


# ------------------------------------------------------------- documents
def _sample_documents() -> Path:
    """Build a Word file and a workbook to exercise the readers and editors."""
    folder = Path(_TEMP_DIR) / "docs"
    folder.mkdir(parents=True, exist_ok=True)

    import docx
    import openpyxl

    document = docx.Document()
    document.add_heading("Case 500XY7 — Contoso Ltd", 0)
    document.add_paragraph("Outlook cannot connect. Error 0x80040115.")
    document.add_paragraph("Status: PENDING")
    document.save(str(folder / "notes.docx"))

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Tickets"
    sheet.append(["Case", "Customer", "Hours", "Status"])
    sheet.append(["500XY7", "Contoso", 4.5, "Open"])
    workbook.save(str(folder / "tickets.xlsx"))
    workbook.close()

    return folder


def test_document_reading():
    print("\nDocuments — reading")
    from app.services import document_readers as readers

    folder = _sample_documents()

    word = readers.read(folder / "notes.docx")
    check("Word document read", word["kind"] == "docx")
    check("Word text extracted", "0x80040115" in word["text"])
    check("Word headings found", len(word["outline"]["headings"]) >= 1)

    excel = readers.read(folder / "tickets.xlsx")
    check("Excel workbook read", excel["kind"] == "xlsx")
    check("Excel sheet named", excel["outline"]["sheets"][0]["name"] == "Tickets")
    check("Excel headers found",
          excel["outline"]["sheets"][0]["headers"][0] == "Case")
    check("Excel values extracted", "Contoso" in excel["text"])

    (folder / "notes.txt").write_text("Plain text notes", encoding="utf-8")
    check("Plain text read", readers.read(folder / "notes.txt")["kind"] == "text")

    for label, path, expected in [
        ("missing file", folder / "nope.docx", "not found"),
        ("unsupported type", folder / "thing.zip", "does not read"),
        ("legacy format", folder / "old.doc", "Word 97"),
    ]:
        if label != "missing file":
            path.write_text("x", encoding="utf-8")
        try:
            readers.read(path)
            check(f"{label} rejected", False)
        except readers.DocumentError as exc:
            check(f"{label} rejected with a clear message", expected in str(exc))


def test_document_editing():
    print("\nDocuments — editing")
    from app.services import document_editors as editors
    from app.services import document_readers as readers

    folder = _sample_documents()
    word = folder / "notes.docx"
    excel = folder / "tickets.xlsx"

    operations = [{"op": "replace_text", "find": "PENDING", "replace": "COMPLETE"}]
    preview = editors.apply(word, operations, dry_run=True)
    check("Dry run reports the change", preview["operations"][0]["occurrences"] == 1)
    check("Dry run writes nothing", "PENDING" in readers.read(word)["text"])
    check("Dry run makes no backup", preview["backup"] is None)

    applied = editors.apply(word, operations)
    check("Edit applied", "COMPLETE" in readers.read(word)["text"])
    check("Original text gone", "PENDING" not in readers.read(word)["text"])
    check("Backup created", applied["backup"] and Path(applied["backup"]).exists())

    editors.apply(excel, [
        {"op": "set_cell", "sheet": "Tickets", "cell": "D2", "value": "Closed"},
        {"op": "append_row", "sheet": "Tickets", "values": ["500ZZ1", "Fabrikam", 2, "Open"]},
        {"op": "set_cell", "sheet": "Tickets", "cell": "E1", "value": "=C2*150"},
    ])
    result = readers.read(excel)
    check("Cell updated", "Closed" in result["text"])
    check("Row appended", "Fabrikam" in result["text"])
    check("Row count grew", result["outline"]["sheets"][0]["rows"] == 3)

    import openpyxl
    workbook = openpyxl.load_workbook(str(excel))
    check("Formula stored as a formula", workbook["Tickets"]["E1"].value == "=C2*150")
    check("Numbers stayed numeric",
          isinstance(workbook["Tickets"].cell(row=3, column=3).value, (int, float)))
    workbook.close()

    for label, target, operations in [
        ("unknown sheet", excel, [{"op": "set_cell", "sheet": "Ghost", "cell": "A1", "value": 1}]),
        ("unknown operation", excel, [{"op": "explode"}]),
        ("paragraph out of range", word, [{"op": "set_paragraph", "index": 999, "text": "x"}]),
        ("no operations", excel, []),
    ]:
        try:
            editors.apply(target, operations)
            check(f"{label} rejected", False)
        except readers.DocumentError:
            check(f"{label} rejected", True)

    # Office holds a ~$ lock file open while a document is open.
    lock = excel.parent / f"~${excel.name}"
    lock.write_text("", encoding="utf-8")
    try:
        editors.apply(excel, [{"op": "set_cell", "cell": "A1", "value": "x"}])
        check("Refuses to edit a file open in Office", False)
    except readers.DocumentError as exc:
        check("Refuses to edit a file open in Office", "open in Office" in str(exc))
    finally:
        lock.unlink()


def test_sharepoint_resolution():
    """A SharePoint URL must resolve to the locally synced file."""
    print("\nDocuments — SharePoint links")
    from app.services import document_service

    synced = Path(_TEMP_DIR) / "sync" / "Contoso Ltd" / "Shared Documents"
    synced.mkdir(parents=True, exist_ok=True)
    (synced / "Q3 Report.xlsx").write_bytes(
        (_sample_documents() / "tickets.xlsx").read_bytes())

    for url, expected in [
        ("https://contoso.sharepoint.com/sites/S/Shared%20Documents/Q3%20Report.xlsx",
         "Q3 Report.xlsx"),
        ("https://contoso.sharepoint.com/:x:/r/sites/S/_layouts/15/Doc.aspx"
         "?sourcedoc=%7Babc%7D&file=Q3%20Report.xlsx", "Q3 Report.xlsx"),
    ]:
        check(f"Filename extracted from {url.split('/')[-1][:24]}",
              document_service.filename_from_url(url) == expected)
        resolved = document_service.resolve_sharepoint(url)
        check("Resolved to the synced file",
              resolved is not None and resolved.name == expected)

    check("Non-document SharePoint page ignored",
          document_service.filename_from_url(
              "https://contoso.sharepoint.com/sites/S/SitePages/Home.aspx") is None)
    check("Unsynced document returns nothing",
          document_service.resolve_sharepoint(
              "https://contoso.sharepoint.com/sites/S/Missing.docx") is None)
    check("SharePoint host recognised",
          document_service.is_sharepoint_url("https://contoso.sharepoint.com/x"))
    check("Ordinary host not treated as SharePoint",
          not document_service.is_sharepoint_url("https://example.com/x.docx"))


def test_document_tracking():
    print("\nDocuments — tracking")
    from app.services.document_service import DocumentService

    folder = _sample_documents()
    db = session()
    service = DocumentService(db)

    record = service.observe(str(folder / "notes.docx"), discovered_by="desktop_watcher")
    check("Document tracked", record.id is not None)
    check("Kind detected", record.kind == "docx")
    check("Case found in the content", record.case_id == "500XY7")
    check("Text extracted", "0x80040115" in (record.text_preview or ""))

    again = service.observe(str(folder / "notes.docx"))
    check("Observing again reuses the same record", again.id == record.id)

    content = service.content(record)
    check("Content served with an outline", "headings" in content["outline"])
    check("Not reported as open in Office", content["open_in_office"] is False)

    indexed = service.index_into_knowledge(record)
    check("Indexed into the knowledge base", indexed["ok"])

    from app.services.rag_service import RAGService
    results = RAGService(db).search("outlook cannot connect error", limit=3)
    check("Document findable by search",
          any("notes.docx" in r["title"] for r in results))
    db.close()


# --------------------------------------------------------------- redaction
def test_redaction():
    print("\nRedaction — secrets and sensitive windows")
    from app.services.redaction import looks_sensitive, redact

    cleaned, fired = redact("my password is Hunter2! and the api key: sk-abc123def456ghi789jkl")
    check("Password removed", "Hunter2" not in cleaned and "password" in fired)
    check("Secret removed", "sk-abc123" not in cleaned)

    cleaned, fired = redact("card 4532 0151 1283 0366 today", redact_pii=True)
    check("Valid card number removed", "4532" not in cleaned)
    cleaned, _ = redact("order 1234 5678 9012 3456 7890")
    check("Non-Luhn long number kept", "1234 5678" in cleaned)

    clean_text = "Outlook error 0x80040115 after the VPN change."
    cleaned, fired = redact(clean_text)
    check("Clean text is untouched", cleaned == clean_text and not fired)

    check("Login window flagged sensitive", looks_sensitive("Sign in - Portal"))
    check("Password manager flagged", looks_sensitive("1Password"))
    check("Ordinary window not flagged", not looks_sensitive("Case 500XY7 - Salesforce"))


def test_activity_capture():
    print("\nActivity capture — storage and guards")
    import os as _os

    _os.environ["ACTIVITY_CAPTURE_ENABLED"] = "true"
    _os.environ["ACTIVITY_EXCLUDED_APPS"] = "notepad"
    from app.core.config import settings
    from app.services.activity_service import ActivityService

    for key in ("ACTIVITY_CAPTURE_ENABLED", "ACTIVITY_EXCLUDED_APPS"):
        object.__setattr__(settings, key,
                           True if key.endswith("ENABLED") else "notepad")

    db = session()
    service = ActivityService(db)

    kept = service.record(kind="keystrokes", application="Teams",
                          window_title="Chat", text="the password is Hunter2 ok")
    check("Frame stored", kept is not None)
    check("Secret redacted before storage", "Hunter2" not in (kept.text or ""))

    dropped = service.record(kind="screenshot", window_title="Bank - Sign in")
    check("Sensitive window dropped", dropped is None)

    excluded = service.record(kind="window", application="notepad.exe",
                              window_title="notepad - secret")
    check("Excluded app dropped", excluded is None)

    result = service.purge(everything=True)
    check("Purge clears everything", result["ok"])


# --------------------------------------------------------------- memory
def test_memory():
    print("\nMemory — store, recall, dedupe")
    from app.services.memory_service import MemoryService

    db = session()
    service = MemoryService(db)
    service.remember("Contoso Outlook fix",
                     "For Contoso, Outlook 0x80040115 after a VPN change was fixed "
                     "by rebuilding the MAPI profile.",
                     memory_type="case_resolution", case_id="500XY7", customer="Contoso")
    service.remember("Northwind prefers email",
                     "Northwind's IT lead prefers email over Teams.",
                     memory_type="customer_fact", customer="Northwind")

    results = service.recall("contoso outlook cannot connect after vpn",
                             case_id="500XY7", customer="Contoso")
    check("Relevant memory recalled", results and "MAPI" in results[0]["content"])
    check("Recall bumps use count", results[0]["use_count"] >= 1)

    before = len(service.list_memories())
    service.remember("Contoso Outlook",
                     "Contoso Outlook error 0x80040115 after VPN fixed by rebuilding "
                     "the MAPI profile.",
                     memory_type="case_resolution", case_id="500XY7", customer="Contoso")
    check("Near-duplicate merged, not added", len(service.list_memories()) == before)

    text = service.recall_text("how to reach northwind")
    check("Recall block is prompt-ready", "Northwind" in text)


# ------------------------------------------------------------- style
def test_style_and_persona():
    print("\nStyle & persona")
    from app.core.config import settings
    from app.services.style_service import StyleService, persona_directive

    db = session()
    style = StyleService(db)
    for sample in [
        "Hey Randy — got it working, mail's flowing again. Lmk if anything pops up. Thanks!",
        "Hi team, quick update: VPN change caused it. Fixed now. Cheers",
        "Thanks for the heads up, I'll take a look this afternoon. Best",
    ]:
        style.add_sample(sample)

    result = style.learn()
    check("Voice learned from samples", result["ok"])
    check("Casual tone detected", result["profile"]["formality"] == "casual")
    check("Greeting picked up", result["profile"]["typical_greeting"] == "hey")

    directive = style.drafting_directive()
    check("Drafting directive produced", "voice" in directive.lower())

    style.add_sample("the login password is Hunter2 use that")
    import json as _json
    last = _json.loads(style._row().samples)[-1]["text"]
    check("Secrets redacted from samples", "Hunter2" not in last)

    object.__setattr__(settings, "PERSONA", "partner")
    check("Partner persona says we/us", "we" in persona_directive().lower())
    object.__setattr__(settings, "PERSONA", "assistant")
    check("Assistant persona says you", "you" in persona_directive().lower())


# ------------------------------------------------------------- tasks
def test_task_parsing_and_scheduling():
    print("\nTasks — parsing and scheduling")
    from datetime import datetime

    from app.services.task_service import (TaskService, compute_next_run,
                                           parse_instruction)

    parsed = parse_instruction("keep the project log updated daily at 9am under my name")
    check("Daily schedule parsed", parsed["schedule"] == "daily")
    check("Time parsed", parsed["at_time"] == "09:00")
    check("Attribution parsed", parsed["attribution"] == "user")
    check("Document-update kind inferred", parsed["kind"] == "document_update")

    parsed = parse_instruction("remind me to call Randy at 3pm")
    check("3pm parsed as 15:00", parsed["at_time"] == "15:00")

    nxt = compute_next_run("daily", "09:00")
    check("Next run is in the future", nxt and nxt > datetime.now())
    check("Manual schedule never fires", compute_next_run("manual") is None)

    db = session()
    task = TaskService(db).create_from_instruction("remind me to update Northwind daily at 9am")
    check("Task created and scheduled", task.next_run is not None)

    # A one-off task must finish after firing, not silently become recurring.
    from datetime import datetime as _dt
    from app.models.task import Task as _Task
    once = _Task(title="call Randy", kind="reminder", schedule="once",
                 at_time="15:00", status="active",
                 next_run=_dt.now())
    db.add(once); db.commit(); db.refresh(once)
    TaskService(db).run(once)
    check("One-off task completes after firing", once.status == "done")
    check("One-off task does not reschedule", once.next_run is None)

    # A weekly task steps a full week, not a day.
    from app.services.task_service import compute_next_run as _next
    first = _next("weekly", "09:00")
    second = _next("weekly", "09:00", after=first)
    check("Weekly reschedule advances seven days",
          (second - first).days == 7)


def test_document_update_task():
    print("\nTasks — autonomous document maintenance")
    import json as _json

    import docx

    from app.models.task import Task
    from app.services import document_readers
    from app.services.task_executors import _find_user_section, execute

    folder = Path(_TEMP_DIR) / "tasklog"
    folder.mkdir(parents=True, exist_ok=True)
    document = docx.Document()
    document.add_heading("Project Log", 0)
    document.add_heading("Team updates", 1)
    document.add_paragraph("2026-08-15 — Kickoff.")
    document.add_heading("My daily log", 1)
    document.add_paragraph("2026-08-16 — Reviewed plan.")
    path = folder / "log.docx"
    document.save(str(path))

    content = document_readers.read(path)
    mine = _find_user_section(content, "mine", "user")
    paragraphs = [p.text for p in docx.Document(str(path)).paragraphs]
    check("'Mine' resolves to the user's section",
          mine is not None and paragraphs[mine - 1] == "My daily log")

    # With a stub LLM, the autonomous write lands under the user's section.
    import app.services.llm_service as llm_module

    original_enabled = llm_module.LLMService.enabled
    original_call = llm_module.LLMService._call_llm
    llm_module.LLMService.enabled = property(lambda self: True)
    llm_module.LLMService._call_llm = lambda self, prompt: "Cut over two servers."
    try:
        db = session()
        task = Task(title="daily log", kind="document_update", autonomous=True,
                    attribution="user", status="active",
                    spec=_json.dumps({"document": str(path), "section": "mine"}))
        db.add(task); db.commit(); db.refresh(task)
        result = execute(db, task)
        check("Autonomous write succeeded", result.get("status") == "active")

        after = [p.text for p in docx.Document(str(path)).paragraphs]
        entry = [p for p in after if "Cut over two servers" in p]
        check("Entry written", bool(entry))
        check("Entry under the user's section",
              after.index(entry[0]) > after.index("My daily log"))
        check("Team section untouched", "Kickoff." in " ".join(after))
    finally:
        llm_module.LLMService.enabled = original_enabled
        llm_module.LLMService._call_llm = original_call


# ------------------------------------------------------------- nudges
def test_nudges():
    print("\nNudges — detection, dedupe, persona")
    from datetime import datetime, timedelta

    from app.core.config import settings
    from app.models.case import Case
    from app.models.enterprise import EnterpriseMessage
    from app.models.event import Event
    from app.services.nudge_service import NudgeService

    db = session()
    db.add(EnterpriseMessage(source="outlook", type="email", external_id="nudge-1",
        sender="randy@company.com", sender_name="Randy", subject="Follow-up",
        urgency="high", handled=False,
        ingested_at=datetime.utcnow() - timedelta(hours=5)))
    db.add(Event(event_type="REMOTE_SESSION_DISCONNECTED", source="agent",
                 case_id="500ZZ9", data={"case_id": "500ZZ9"},
                 created_at=datetime.utcnow() - timedelta(hours=2)))
    db.add(Case(case_id="500ZZ9", system="Salesforce", customer="Fabrikam",
                title="Migration", status="open"))
    db.commit()

    service = NudgeService(db)
    object.__setattr__(settings, "PERSONA", "partner")
    first = service.scan()
    check("Nudges raised", first["raised"] >= 2)

    nudges = {n.kind: n for n in service.open_nudges()}
    check("Unanswered-mail nudge raised", "unanswered_email" in nudges)
    check("Case-not-updated nudge raised", "case_not_updated" in nudges)
    check("Partner voice used", "we" in nudges["case_not_updated"].body.lower())

    second = service.scan()
    check("Re-scan does not duplicate", second["raised"] == 0)


# ------------------------------------------------------ copilot bridge
def test_copilot_bridge():
    """Publishing, command execution, and the boundaries that must hold."""
    print("\nCopilot bridge")
    import json as _json

    from app.core.config import settings
    from app.services.copilot_bridge import ALLOWED_COMMANDS, CopilotBridge

    object.__setattr__(settings, "COPILOT_BRIDGE_ENABLED", True)
    object.__setattr__(settings, "COPILOT_COMMAND_MODE", "auto")

    folder = Path(_TEMP_DIR) / "copilot"
    db = session()
    bridge = CopilotBridge(db)

    published = bridge.publish()
    check("Publishes the shared files", published["ok"])
    check("Context published", (folder / "context.json").exists())
    check("Memory published", (folder / "memory.json").exists())
    check("Style published", (folder / "style.json").exists())

    context = _json.loads((folder / "context.json").read_text())
    check("Context carries a freshness stamp", "generated_at" in context)
    check("Context includes suggestions", "suggestions" in context)

    memory = _json.loads((folder / "memory.json").read_text())
    check("Memory payload has no raw activity",
          not any(key in memory for key in ("screenshots", "keystrokes", "activity")))

    # A permitted command runs.
    commands = folder / "commands"
    commands.mkdir(parents=True, exist_ok=True)
    good = commands / "cmd-good.json"
    good.write_text(_json.dumps({"action": "get_context"}), encoding="utf-8")
    os.utime(good, (0, 0))
    result = bridge.drain_commands()
    check("Permitted command executed", result["executed"] == 1)
    check("Command file archived", not good.exists())

    results = list((folder / "results").glob("*.json"))
    check("Result written back for the agent", bool(results))

    # A command outside the whitelist is refused, not executed.
    bad = commands / "cmd-bad.json"
    bad.write_text(_json.dumps({"action": "send_email", "to": "ceo@company.com"}),
                   encoding="utf-8")
    os.utime(bad, (0, 0))
    bridge.drain_commands()
    refusal = _json.loads(
        sorted((folder / "results").glob("*cmd-bad*"))[0].read_text())
    check("Non-permitted command refused", refusal["ok"] is False)
    check("Refusal explains why", "not-permitted" in refusal["error"]
          or "Unknown" in refusal["error"])
    check("Sending is not a permitted action", "send_email" not in ALLOWED_COMMANDS)

    # Approval mode stages changes instead of doing them.
    object.__setattr__(settings, "COPILOT_COMMAND_MODE", "approve")
    staged = commands / "cmd-change.json"
    staged.write_text(_json.dumps({
        "action": "append_document", "name": "nope.docx", "text": "hi"}),
        encoding="utf-8")
    os.utime(staged, (0, 0))
    bridge.drain_commands()
    outcome = _json.loads(
        sorted((folder / "results").glob("*cmd-change*"))[0].read_text())
    check("Changes wait for approval in approve mode", outcome.get("staged") is True)

    status = bridge.status()
    check("Status reports the folder", status["enabled"] and status["ok"])
    db.close()


def test_copilot_guide():
    """The pasteable instructions must actually describe the real contract."""
    print("\nCopilot guide")
    from app.services.copilot_bridge import ALLOWED_COMMANDS
    from app.services.copilot_guide import guide

    payload = guide()
    check("Guide has steps", len(payload["steps"]) >= 5)
    check("OneDrive is the required tool",
          any(t["required"] and "OneDrive" in t["name"] for t in payload["tools"]))

    instructions = payload["instructions"]
    check("Instructions mention context.json", "context.json" in instructions)
    check("Instructions mention style.json", "style.json" in instructions)
    check("Instructions forbid sending without approval",
          "without showing the user the draft" in instructions)
    check("Instructions tell it not to invent desktop state",
          "Never invent desktop state" in instructions)

    # Every command the instructions advertise must really be permitted, or the
    # agent will be told to do things Cerebro refuses.
    advertised = {name for name in ALLOWED_COMMANDS if f'"{name}"' in instructions}
    check("Every advertised command is permitted",
          advertised and advertised.issubset(set(ALLOWED_COMMANDS)))
    check("All permitted commands are documented",
          set(ALLOWED_COMMANDS).issubset(advertised))


def test_copilot_approval_flow():
    """A change staged in approve mode must actually run when approved."""
    print("\nCopilot approval flow")
    import json as _json

    import docx

    from app.core.config import settings
    from app.models.nudge import Nudge
    from app.services.copilot_bridge import CopilotBridge
    from app.services.document_service import DocumentService
    from app.services.nudge_service import NudgeService

    object.__setattr__(settings, "COPILOT_BRIDGE_ENABLED", True)
    object.__setattr__(settings, "COPILOT_COMMAND_MODE", "approve")

    folder = Path(_TEMP_DIR) / "approve"
    (folder / "commands").mkdir(parents=True, exist_ok=True)
    object.__setattr__(settings, "COPILOT_BRIDGE_DIR", str(folder))

    # A tracked document with a section that is the user's.
    document = docx.Document()
    document.add_heading("Log", 0)
    document.add_heading("My updates", 1)
    document.add_paragraph("2026-08-16 — started.")
    path = folder / "log.docx"
    document.save(str(path))

    db = session()
    DocumentService(db).observe(str(path))

    # Stub the LLM so the append can produce text.
    import app.services.llm_service as llm_module

    original_enabled = llm_module.LLMService.enabled
    original_call = llm_module.LLMService._call_llm
    llm_module.LLMService.enabled = property(lambda self: True)
    llm_module.LLMService._call_llm = lambda self, prompt: "reviewed the migration."
    try:
        bridge = CopilotBridge(db)
        command = {"action": "append_document", "name": "log.docx",
                   "section": "mine", "text": "add a status line"}

        # Stage via the real path: write a command file and drain.
        cmd = folder / "commands" / "cmd-approve.json"
        cmd.write_text(_json.dumps(command), encoding="utf-8")
        os.utime(cmd, (0, 0))
        bridge.drain_commands()

        # The shared test DB may hold copilot_request nudges from earlier tests;
        # take the most recent, which is the one we just staged.
        nudge = (db.query(Nudge)
                 .filter(Nudge.kind == "copilot_request")
                 .order_by(Nudge.id.desc()).first())
        check("Change staged as a nudge", nudge is not None)

        # Approving the nudge must actually perform the append.
        from app.api.tasks import act

        act(nudge.id, db)
        after = [p.text for p in docx.Document(str(path)).paragraphs]
        check("Approved command actually ran",
              any("reviewed the migration" in p for p in after))
        check("Entry landed under the user's section",
              after.index([p for p in after if "reviewed the migration" in p][0])
              > after.index("My updates"))

        # A duplicate stage does not silently vanish.
        cmd2 = folder / "commands" / "cmd-approve-2.json"
        cmd2.write_text(_json.dumps(command), encoding="utf-8")
        os.utime(cmd2, (0, 0))
        bridge.drain_commands()
        result2 = _json.loads(
            sorted((folder / "results").glob("*cmd-approve-2*"))[0].read_text())
        check("Duplicate stage reports it is already pending",
              result2.get("duplicate") is True)
    finally:
        llm_module.LLMService.enabled = original_enabled
        llm_module.LLMService._call_llm = original_call
    db.close()


def test_copilot_memory_redaction():
    """Memory published to the cloud folder must be redacted."""
    print("\nCopilot memory redaction")
    import json as _json

    from app.core.config import settings
    from app.services.copilot_bridge import CopilotBridge
    from app.services.memory_service import MemoryService

    folder = Path(_TEMP_DIR) / "redact-pub"
    folder.mkdir(parents=True, exist_ok=True)
    object.__setattr__(settings, "COPILOT_BRIDGE_ENABLED", True)
    object.__setattr__(settings, "COPILOT_BRIDGE_DIR", str(folder))

    db = session()
    # A memory that (wrongly) captured a secret in its content.
    MemoryService(db).remember(
        "Server access", "The admin password is Hunter2 for the Contoso box.",
        memory_type="fact")

    CopilotBridge(db).publish()
    published = _json.loads((folder / "memory.json").read_text())
    blob = _json.dumps(published)
    check("Secret redacted before publishing", "Hunter2" not in blob)
    check("Memory still published", published["count"] >= 1)
    db.close()


# ------------------------------------------------------------- regressions
def test_embedding_signature_matches_vector():
    """A fallback vector must be labelled with the space it actually belongs to."""
    print("\nEmbedding signatures")
    vector, produced = embeddings.embed_with_signature("outlook exchange error")
    check("Signature matches the vector length",
          produced.endswith(f":{len(vector)}"))
    check("Local provider reports the local signature",
          produced == embeddings.local_signature())


def test_schema_upgrade():
    """A database created before the embedding columns existed must still work."""
    print("\nSchema upgrade")
    import sqlite3
    from sqlalchemy import create_engine, inspect

    legacy = Path(_TEMP_DIR) / "legacy.db"
    connection = sqlite3.connect(legacy)
    connection.execute(
        "CREATE TABLE documents (id INTEGER PRIMARY KEY, source VARCHAR, title VARCHAR, "
        "content TEXT, url VARCHAR, vector_id VARCHAR, tags VARCHAR, indexed BOOLEAN, "
        "created_at DATETIME, updated_at DATETIME)")
    connection.commit()
    connection.close()

    import app.core.database as database

    original = database.engine
    try:
        database.engine = create_engine(f"sqlite:///{legacy.as_posix()}",
                                        connect_args={"check_same_thread": False})
        database.init_db()
        columns = {c["name"] for c in inspect(database.engine).get_columns("documents")}
        check("Missing embedding column is added", "embedding" in columns)
        check("Missing signature column is added", "embedding_signature" in columns)
    finally:
        database.engine = original


def test_customer_from_title():
    """A two-part title carries no customer — guessing files the case number."""
    print("\nTitle parsing")
    detected = EventDetector.detect_crm_event(
        "https://acme.lightning.force.com/lightning/r/Case/5008d00000ABCDEfgh/view",
        "Case 00001234 | Contoso Ltd | Salesforce")
    check("Three-part title yields the customer",
          detected and detected[1]["customer"] == "Contoso Ltd")

    detected = EventDetector.detect_crm_event(
        "https://acme.lightning.force.com/lightning/r/Case/5008d00000ABCDEfgh/view",
        "Case 00001234 | Salesforce")
    check("Two-part title yields no customer",
          detected and detected[1]["customer"] is None)


def test_database_password_masking():
    """The settings API must never hand a database password to the browser."""
    print("\nCredential masking")
    from app.core import settings_store

    url = "postgresql+psycopg://cerebro:s3cret@db.internal:5432/cerebro"
    masked = settings_store.mask_url_password(url)
    check("Password is masked", "s3cret" not in masked)
    check("Rest of the URL survives", "db.internal:5432/cerebro" in masked)
    check("Masked value round-trips back to the real password",
          settings_store.restore_url_password(masked, url) == url)
    check("SQLite URLs are untouched",
          settings_store.mask_url_password("sqlite:///./data/cerebro.db")
          == "sqlite:///./data/cerebro.db")


# ----------------------------------------------------------- settings store
def test_env_round_trip():
    """Windows paths and multi-line settings must survive a write/read cycle."""
    print("\nSettings — .env encoding")
    import tempfile as _tempfile

    from dotenv import dotenv_values

    from app.core.settings_store import _encode

    folder = Path(_tempfile.mkdtemp(dir=_TEMP_DIR))
    cases = [
        ("windows path", r"C:\Users\you\Contoso"),
        # \n and \t inside a quoted value are escape sequences to dotenv.
        ("path with escape-shaped segments", r"C:\notes\team"),
        ("multi-line folders", "C:\\a\\b\nC:\\c\\d"),
        ("multi-line domains", "mybank.com\npayroll.company.com"),
        ("value with a hash", "value # not a comment"),
        ("value with quotes", 'say "hello"'),
        ("padded value", "  padded  "),
    ]
    for label, value in cases:
        target = folder / "t.env"
        target.write_text(f"K={_encode(value)}\nNEXT=ok\n", encoding="utf-8")
        parsed = dotenv_values(target)
        check(f"{label} round-trips",
              parsed.get("K") == value and parsed.get("NEXT") == "ok"
              and len(parsed) == 2)


def test_urgency_word_boundaries():
    """Urgency must key on words, not substrings."""
    print("\nEnterprise bridge — urgency precision")
    from app.services.enterprise_service import assess_urgency

    for text, expected in [
        ("Please download the report", "normal"),
        ("Shutdown window is Saturday", "normal"),
        ("Countdown to launch", "normal"),
        ("The site is down", "high"),
        ("This is escalating fast", "high"),
        ("Escalated to tier 3", "high"),
        ("Server outage in progress", "high"),
        ("Can you review by EOD", "medium"),
    ]:
        check(f"{text!r} → {expected}", assess_urgency(text, "", "normal")[0] == expected)


def test_empty_batch_file():
    """An empty batch must be archived, not retried forever."""
    print("\nEnterprise bridge — empty batch")
    from app.services import enterprise_service

    inbox = Path(_TEMP_DIR) / "inbox"
    inbox.mkdir(parents=True, exist_ok=True)
    empty = inbox / "outlook-empty.json"
    empty.write_text("[]", encoding="utf-8")
    os.utime(empty, (0, 0))

    db = session()
    result = enterprise_service.drain_inbox(db)
    check("Empty batch did not raise", result["ok"])
    check("Empty batch archived, not left to retry", not empty.exists())
    db.close()


def test_office_lock_detection():
    """Office truncates lock-file names for long documents."""
    print("\nDocuments — open-in-Office guard")
    from app.services.document_readers import is_open_in_office

    folder = Path(_TEMP_DIR) / "locks"
    folder.mkdir(parents=True, exist_ok=True)

    short = folder / "notes.docx"
    short.write_text("x", encoding="utf-8")
    check("Unlocked file reported as closed", not is_open_in_office(short))

    (folder / "~$notes.docx").write_text("", encoding="utf-8")
    check("Exact lock name detected", is_open_in_office(short))

    # Office drops leading characters when the base name is long.
    long_name = folder / "Case notes 2026.docx"
    long_name.write_text("x", encoding="utf-8")
    check("Long name reported as closed before locking",
          not is_open_in_office(long_name))
    (folder / "~$se notes 2026.docx").write_text("", encoding="utf-8")
    check("Truncated lock name detected", is_open_in_office(long_name))


def test_spreadsheet_value_coercion():
    """Identifier-shaped strings must not be turned into numbers."""
    print("\nDocuments — cell value coercion")
    from app.services.document_editors import _coerce

    for value, expected in [
        ("00123", "00123"),          # case number, not one hundred and twenty three
        ("+441234567", "+441234567"),  # phone number
        ("-0012", "-0012"),
        ("123", 123),
        ("4.5", 4.5),
        ("=A1*2", "=A1*2"),
        ("Contoso", "Contoso"),
    ]:
        check(f"{value!r} → {expected!r}", _coerce(value) == expected)


def test_failed_edit_leaves_no_litter():
    """A rejected edit must not leave a backup beside the user's document."""
    print("\nDocuments — failed edit cleanup")
    from app.services import document_editors as editors
    from app.services.document_readers import DocumentError

    folder = _sample_documents()
    excel = folder / "tickets.xlsx"
    before = set(folder.glob("*cerebro-backup*"))

    try:
        editors.apply(excel, [{"op": "set_cell", "sheet": "Ghost", "cell": "A1", "value": 1}])
    except DocumentError:
        pass

    check("No backup left behind by a rejected edit",
          set(folder.glob("*cerebro-backup*")) == before)


def test_settings_store():
    print("\nSettings")
    from app.core import settings_store

    described = settings_store.describe()
    check("Every group has fields",
          all(any(f["group"] == g["id"] for f in described["fields"])
              for g in described["groups"]))

    secrets = [f for f in described["fields"] if f["secret"]]
    check("Secrets are masked, never returned",
          all(f["value"] in ("", settings_store.SECRET_MASK) for f in secrets))

    result = settings_store.update({"PORT": "not-a-number"})
    check("Invalid values are rejected", result["ok"] is False)


# -------------------------------------------------------------------- main
def main() -> int:
    print("Running Cerebro tests…")
    init_db()

    for suite in (test_event_detector, test_context_engine, test_event_flow,
                  test_embeddings, test_knowledge_search, test_llm_disabled,
                  test_enterprise_normalisation, test_enterprise_ingest_and_reply,
                  test_document_reading, test_document_editing,
                  test_sharepoint_resolution, test_document_tracking,
                  test_embedding_signature_matches_vector, test_schema_upgrade,
                  test_customer_from_title, test_database_password_masking,
                  test_env_round_trip, test_urgency_word_boundaries,
                  test_empty_batch_file, test_office_lock_detection,
                  test_spreadsheet_value_coercion, test_failed_edit_leaves_no_litter,
                  test_redaction, test_activity_capture, test_memory,
                  test_style_and_persona, test_task_parsing_and_scheduling,
                  test_document_update_task, test_nudges,
                  test_copilot_bridge, test_copilot_guide,
                  test_copilot_approval_flow, test_copilot_memory_redaction,
                  test_settings_store):
        try:
            suite()
        except Exception as exc:  # a crashing suite is a failure, not a stack trace
            FAILED.append(f"{suite.__name__} raised {type(exc).__name__}: {exc}")
            print(f"  ✗ {suite.__name__} raised {type(exc).__name__}: {exc}")

    print(f"\n{len(PASSED)} passed, {len(FAILED)} failed")
    if FAILED:
        print("\nFailures:")
        for failure in FAILED:
            print(f"  - {failure}")
    return 1 if FAILED else 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    finally:
        shutil.rmtree(_TEMP_DIR, ignore_errors=True)
